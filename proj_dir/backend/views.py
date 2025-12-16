from django.shortcuts import render
from django.contrib.auth import login
from docx.styles.style import ParagraphStyle
from reportlab.lib.pagesizes import A4
import copy
from .forms import UserRegisterForm
import os
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from .models import Analysis
from .utils import extract_pdf_text, extract_docx_text, extract_txt_text
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document as DocxDocument
import requests
import json
from dotenv import load_dotenv
from django import template

from .models import Analysis
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404


from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


register = template.Library()

@register.filter
def severity_color(severity):
    if severity >= 4:
        return 'danger'
    elif severity >= 3:
        return 'warning'
    return 'info'
# Загружаем переменные окружения из .env файла
load_dotenv()
HF_TOKEN = os.getenv('HF_TOKEN', '')

def register(request):
    if request.method == 'POST':
        form = UserRegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            messages.success(request, f'Добро пожаловать, {username}! Ваш аккаунт успешно создан.')
            login(request, user)  # Автоматический вход пользователя
            return redirect('upload')  # Перенаправление на страницу загрузки
    else:
        form = UserRegisterForm()
    return render(request, 'registration/register.html', {'form': form})


import os
import json
import requests
from django.conf import settings


def analyze_document_with_ai(text):
    """Анализирует текст документа с помощью нейросети - УЛУЧШЕННАЯ ВЕРСИЯ"""
    # Получаем токен из настроек или переменных окружения
    HF_TOKEN = getattr(settings, 'HF_TOKEN', os.getenv('HF_TOKEN', ''))

    if not HF_TOKEN:
        return {
            "issues": [{
                "metric": 0,
                "text": "Ошибка конфигурации",
                "severity": 5,
                "explanation": "Не указан токен Hugging Face API"
            }],
            "suggestions": [
                "Настройте переменную окружения HF_TOKEN или добавьте ее в settings.py",
                "Для тестирования можно использовать временный токен в настройках"
            ],
            "risk_score": 0.9
        }

    headers = {
        "Authorization": f"Bearer {HF_TOKEN}",
        "Content-Type": "application/json"
    }

    # ДЕТАЛЬНЫЙ ПРОМПТ С КОНКРЕТНЫМИ ИНСТРУКЦИЯМИ
    prompt = (
        "ТЫ - ЭКСПЕРТ ПО ЮРИДИЧЕСКОМУ И БИЗНЕС-АНАЛИЗУ ДОКУМЕНТОВ.\n"
        "Проанализируй ПРИКРЕПЛЕННЫЙ ДОКУМЕНТ и найди проблемы по 8 критериям ниже.\n\n"

        "КРИТЕРИИ АНАЛИЗА:\n"
        "1. СКРЫТЫЕ КОМИССИИ И ДОП. ПЛАТЕЖИ: Найди упоминания платежей, которые не указаны явно в основной стоимости или описаны расплывчато\n"
        "2. СТРАННЫЕ/ЗАВЫШЕННЫЕ ЦЕНЫ: Найди цены, которые явно не соответствуют рыночным или не имеют обоснования\n"
        "3. НЕПОЛНОТА ДАННЫХ: Проверь наличие контактов, реквизитов, четких сроков и конкретных объемов работ\n"
        "4. РИСКОВАННЫЕ/ОДНОСТОРОННИЕ УСЛОВИЯ: Найди условия, выгодные только одной стороне или создающие необоснованные риски\n"
        "5. ОТСУТСТВУЮЩИЕ ГАРАНТИИ: Проверь наличие конкретных гарантий качества, сроков, результатов\n"
        "6. НЕЧЕТКИЕ ФОРМУЛИРОВКИ: Найди фразы типа \"по договоренности\", \"возможны изменения\", \"примерно\" без конкретики\n"
        "7. ЮРИДИЧЕСКИ ОПАСНЫЕ ФРАЗЫ: Найди упоминания штрафов без обоснования, отказов от ответственности, неограниченной ответственности\n"
        "8. НЕПРОЗРАЧНОСТЬ: Общая оценка ясности и понятности условий для обеих сторон\n\n"

        "ИНСТРУКЦИИ ДЛЯ ФОРМИРОВАНИЯ ОТВЕТА:\n"
        "1. ПРОАНАЛИЗИРУЙ КАЖДОЕ ПРЕДЛОЖЕНИЕ ДОКУМЕНТА ТЩАТЕЛЬНО\n"
        "2. ДЛЯ КАЖДОЙ НАЙДЕННОЙ ПРОБЛЕМЫ УКАЖИ:\n"
        "   - metric: номер критерия (1-8)\n"
        "   - text: ТОЧНАЯ ЦИТАТА из документа (минимум 3 слова, максимум 1 предложение)\n"
        "   - severity: 1-5 (5 = нарушение закона/крайне рискованно)\n"
        "   - explanation: 1 предложение - ПОЧЕМУ это проблема\n"
        "3. В suggestions дай КОНКРЕТНЫЕ ШАГИ ПО ИСПРАВЛЕНИЮ:\n"
        "   - Каждая рекомендация должна начинаться с глагола (\"Добавьте\", \"Уберите\", \"Замените\", \"Уточните\")\n"
        "   - Укажи ТОЧНО, что и где нужно изменить\n"
        "   - Добавь пример правильной формулировки\n"
        "4. risk_score: рассчитай объективно 0.0-1.0 на основе количества и серьезности проблем\n\n"

        "ПРИМЕРЫ КОНКРЕТНЫХ РЕКОМЕНДАЦИЙ:\n"
        "\"Вместо 'дополнительные платежи по договоренности' укажите конкретные виды платежей и их сумму в разделе 2.3: 'Дополнительно оплачиваются: транспортные расходы - 15 000 руб., командировочные - 10 000 руб.'\"\n"
        "\"Добавьте в раздел 7.1 конкретные гарантийные обязательства: 'Гарантийный срок - 12 месяцев с даты подписания акта, в течение которого Исполнитель обязуется безвозмездно устранить все выявленные недостатки'\"\n"
        "\"Уберите фразу 'Заказчик отказывается от всех претензий' из пункта 9.1 - это нарушает ст. 16 Закона о защите прав потребителей. Замените на: 'Заказчик имеет право предъявить претензии в установленном законом порядке'\"\n"
        "\"Замените 'по договоренности' в пункте 4.2 на 'в течение 3 рабочих дней с момента получения письменного запроса'\"\n"
        "\"Укажите конкретные сроки оплаты в разделе 2.2: 'Оплата производится в течение 10 рабочих дней с даты выставления счета'\"\n\n"

        "ПРИМЕРЫ ОБЩИХ (НЕПРАВИЛЬНЫХ) РЕКОМЕНДАЦИЙ:\n"
        "\"Улучшите документ\" (слишком общее)\n"
        "\"Добавьте больше деталей\" (не конкретно)\n"
        "\"Сделайте условия более справедливыми\" (расплывчато)\n\n"

        "ВАЖНЫЕ ПРАВИЛА:\n"
        "- Если проблема не найдена по какому-то критерию - НЕ УПОМИНАЙ этот критерий в ответе\n"
        "- ВСЕГДА ВОЗВРАЩАЙ ЧИСТЫЙ JSON БЕЗ ДОПОЛНИТЕЛЬНОГО ТЕКСТА\n"
        "- УБЕДИСЬ, ЧТО ОТВЕТ СООТВЕТСТВУЕТ ФОРМАТУ JSON\n"
        "- ЕСЛИ НЕ УВЕРЕН В ЧЕМ-ТО - НЕ ВКЛЮЧАЙ ЭТО В ОТВЕТ\n\n"

        "ТЕКСТ ДОКУМЕНТА ДЛЯ АНАЛИЗА:\n"
        "=== НАЧАЛО ДОКУМЕНТА ===\n"
        f"{text}\n"
        "=== КОНЕЦ ДОКУМЕНТА ===\n\n"

        "ВЫВОД ДОЛЖЕН БЫТЬ В СТРОГОМ JSON ФОРМАТЕ БЕЗ ДОПОЛНИТЕЛЬНОГО ТЕКСТА:\n"
        "{\n"
        "  \"issues\": [\n"
        "    {\n"
        "      \"metric\": 1,\n"
        "      \"text\": \"точная цитата\",\n"
        "      \"severity\": 4,\n"
        "      \"explanation\": \"конкретное объяснение проблемы\"\n"
        "    }\n"
        "  ],\n"
        "  \"suggestions\": [\n"
        "    \"Конкретная рекомендация с примером исправления\",\n"
        "    \"Еще одна конкретная рекомендация\"\n"
        "  ],\n"
        "  \"risk_score\": 0.65\n"
        "}\n"
    )

    payload = {
        "model": "Qwen/Qwen2.5-7B-Instruct",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1500,
        "temperature": 0.2,  # Более детерминированный ответ
        "top_p": 0.85
    }

    try:
        # Отправляем запрос к API
        response = requests.post(
            "https://router.huggingface.co/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=90  # Увеличиваем таймаут для сложного анализа
        )

        # Проверяем статус ответа
        if response.status_code != 200:
            return {
                "issues": [{
                    "metric": 0,
                    "text": f"Ошибка API: {response.status_code}",
                    "severity": 5,
                    "explanation": f"Сервер вернул ошибку: {response.text[:200]}"
                }],
                "suggestions": [
                    f"Проверьте токен Hugging Face API (текущий статус: {response.status_code})",
                    "Убедитесь, что токен действителен и имеет доступ к модели",
                    "Попробуйте повторить запрос позже"
                ],
                "risk_score": 0.8
            }

        # Парсим ответ
        result = response.json()
        if "choices" not in result or len(result["choices"]) == 0:
            return {
                "issues": [{
                    "metric": 0,
                    "text": "Некорректный формат ответа",
                    "severity": 5,
                    "explanation": "API вернул ответ без choices массива"
                }],
                "suggestions": [
                    "Проверьте логи сервера для отладки",
                    "Убедитесь, что используете правильный формат запроса к API"
                ],
                "risk_score": 0.7
            }

        # Извлекаем текст ответа
        response_text = result["choices"][0]["message"]["content"]

        # Очищаем ответ, оставляя только JSON
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1

        if json_start == -1 or json_end == 0:
            # Пробуем найти JSON в другом формате
            json_start = response_text.lower().find('```json')
            if json_start != -1:
                json_start += 7
                json_end = response_text.rfind('```')

            if json_start == -1 or json_end == -1:
                return {
                    "issues": [{
                        "metric": 0,
                        "text": "Некорректный формат ответа",
                        "severity": 5,
                        "explanation": "Нейросеть не вернула ответ в формате JSON"
                    }],
                    "suggestions": [
                        "Попробуйте уменьшить длину анализируемого текста",
                        "Проверьте, что документ не содержит специальных символов",
                        "Обратитесь к администратору системы для отладки"
                    ],
                    "risk_score": 0.75
                }

        # Извлекаем чистый JSON
        json_str = response_text[json_start:json_end].strip()

        try:
            # Парсим JSON
            parsed_result = json.loads(json_str)

            # Валидация структуры ответа
            if not isinstance(parsed_result.get("issues"), list):
                parsed_result["issues"] = []

            if not isinstance(parsed_result.get("suggestions"), list):
                parsed_result["suggestions"] = []

            risk_score = parsed_result.get("risk_score", 0.5)
            if not isinstance(risk_score, (int, float)) or risk_score < 0 or risk_score > 1:
                parsed_result["risk_score"] = 0.5

            # Фильтруем некорректные записи
            valid_issues = []
            for issue in parsed_result["issues"]:
                if isinstance(issue, dict) and all(k in issue for k in ["metric", "text", "severity", "explanation"]):
                    if isinstance(issue["text"], str) and len(issue["text"].strip()) >= 3:
                        valid_issues.append(issue)

            parsed_result["issues"] = valid_issues

            return parsed_result

        except json.JSONDecodeError as e:
            return {
                "issues": [{
                    "metric": 0,
                    "text": "Ошибка парсинга JSON",
                    "severity": 5,
                    "explanation": f"Некорректный синтаксис JSON: {str(e)}"
                }],
                "suggestions": [
                    "Повторите анализ",
                    "Попробуйте сократить объем анализируемого текста",
                    "Свяжитесь с администратором для отладки интеграции"
                ],
                "risk_score": 0.7
            }

    except requests.exceptions.Timeout:
        return {
            "issues": [{
                "metric": 0,
                "text": "Таймаут запроса",
                "severity": 5,
                "explanation": "Анализ занял слишком много времени (более 90 секунд)"
            }],
            "suggestions": [
                "Попробуйте проанализировать документ меньшего размера",
                "Повторите попытку позже, когда сервер будет менее загружен",
                "Разделите большой документ на части и проанализируйте отдельно"
            ],
            "risk_score": 0.6
        }
    except requests.exceptions.RequestException as e:
        return {
            "issues": [{
                "metric": 0,
                "text": "Ошибка подключения к API",
                "severity": 5,
                "explanation": f"Невозможно подключиться к серверу: {str(e)}"
            }],
            "suggestions": [
                "Проверьте подключение к интернету",
                "Убедитесь, что сервер API доступен",
                "Проверьте корректность URL эндпоинта API"
            ],
            "risk_score": 0.8
        }
    except Exception as e:
        return {
            "issues": [{
                "metric": 0,
                "text": "Системная ошибка",
                "severity": 5,
                "explanation": f"Неожиданная ошибка: {str(e)}"
            }],
            "suggestions": [
                "Повторите попытку",
                "Обратитесь к администратору системы",
                "Проверьте логи приложения для детальной информации"
            ],
            "risk_score": 0.9
        }


@login_required
def delete_account(request):
    if request.method == 'POST':
        # Удаляем все данные пользователя
        request.user.delete()
        messages.success(request, "Ваш аккаунт был успешно удален.")
        return redirect('home')
    return render(request, 'account/delete_account.html')

@login_required
def upload_document(request):
    if request.method == 'POST' and request.FILES.get('document'):
        file = request.FILES['document']
        file_type = file.name.split('.')[-1].lower()

        # Извлечение текста
        if file_type == 'pdf':
            text = extract_pdf_text(file)
        elif file_type == 'docx':
            text = extract_docx_text(file)
        else:
            text = extract_txt_text(file)

        # Анализ текста с помощью нейросети
        api_result = analyze_document_with_ai(text)

        # Сохранение результата
        analysis = Analysis.objects.create(
            user=request.user,
            document=file,
            text=text,
            result=api_result
        )

        return redirect('report', analysis_id=analysis.id)

    return render(request, 'upload/upload.html')


@login_required
def report(request, analysis_id):
    analysis = Analysis.objects.get(id=analysis_id, user=request.user)
    return render(request, 'report/report.html', {'analysis': analysis})


@login_required
def history(request):
    analyses = Analysis.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'history/history.html', {'analyses': analyses})

@login_required
def download_pdf(request, analysis_id):
    analysis = get_object_or_404(Analysis, id=analysis_id, user=request.user)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="report_{analysis.id}.pdf"'

    # РЕГИСТРИРУЕМ ШРИФТ С ПОДДЕРЖКОЙ КИРИЛЛИЦЫ
    try:
        # Пытаемся использовать системный шрифт DejaVu Sans
        pdfmetrics.registerFont(TTFont('DejaVu', 'DejaVuSans.ttf'))
    except:
        try:
            # Если не получается, используем Arial (для Windows)
            pdfmetrics.registerFont(TTFont('DejaVu', 'Arial.ttf'))
        except:
            # Фallback на шрифт с минимальной поддержкой
            pdfmetrics.registerFont(TTFont('DejaVu', 'Helvetica'))

    # Создаем документ
    doc = SimpleDocTemplate(
        response,
        pagesize=A4,
        rightMargin=30,
        leftMargin=30,
        topMargin=30,
        bottomMargin=30
    )

    # Создаем базовые стили
    styles = getSampleStyleSheet()

    # Настраиваем шрифт для всех стилей
    for style in styles.byName.values():
        style.fontName = 'DejaVu'
        style.fontSize = 10
        style.leading = 12
        style.spaceBefore = 0
        style.spaceAfter = 0

    # Создаем элементы документа
    story = []

    # Заголовок
    story.append(Paragraph("Отчет по анализу документа", styles['Normal']))
    story.append(Spacer(1, 6))

    # Данные о пользователе и дате
    story.append(Paragraph(f"Пользователь: {analysis.user.username}", styles['Normal']))
    story.append(Paragraph(f"Дата анализа: {analysis.created_at.strftime('%d.%m.%Y %H:%M')}", styles['Normal']))
    story.append(Paragraph(f"Имя файла: {analysis.document.name}", styles['Normal']))
    story.append(Spacer(1, 6))

    # Текст документа
    story.append(Paragraph("Проанализированный текст:", styles['Normal']))
    story.append(Spacer(1, 3))

    document_text = analysis.text[:2000]
    if len(analysis.text) > 2000:
        document_text += "... (текст обрезан для сохранения формата PDF)"

    story.append(Paragraph(document_text, styles['Normal']))
    story.append(Spacer(1, 6))

    # Проблемы
    story.append(Paragraph("Выявленные проблемы:", styles['Normal']))
    story.append(Spacer(1, 3))

    if hasattr(analysis.result, 'get') and analysis.result.get('issues'):
        for issue in analysis.result['issues']:
            issue_text = issue.get('text', '')
            explanation = issue.get('explanation', '')
            story.append(Paragraph(f"• {issue_text}", styles['Normal']))
            if explanation:
                story.append(Paragraph(f"  {explanation}", styles['Normal']))
    else:
        story.append(Paragraph("Проблемы не обнаружены.", styles['Normal']))

    story.append(Spacer(1, 6))
    story.append(Paragraph("Рекомендации:", styles['Normal']))
    story.append(Spacer(1, 3))

    if hasattr(analysis.result, 'get') and analysis.result.get('suggestions'):
        for suggestion in analysis.result['suggestions']:
            story.append(Paragraph(f"• {suggestion}", styles['Normal']))
    else:
        story.append(Paragraph("Рекомендации отсутствуют.", styles['Normal']))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Отчет сгенерирован системой анализа бизнес-предложений", styles['Normal']))

    doc.build(story)
    return response

@login_required
def clear_history(request):
    if request.method == 'POST':
        # Удаляем все анализы пользователя
        Analysis.objects.filter(user=request.user).delete()
        messages.success(request, "История анализов была очищена.")
        return redirect('history')
    return render(request, 'history/clear_history.html')

@login_required
def download_txt(request, analysis_id):
    analysis = Analysis.objects.get(id=analysis_id, user=request.user)

    # Формируем текстовый отчет
    report = f"ОТЧЕТ ПО АНАЛИЗУ ДОКУМЕНТА\n\n"
    report += f"Дата анализа: {analysis.created_at.strftime('%d.%m.%Y %H:%M')}\n"
    report += f"Анализируемый документ: {analysis.document.name}\n\n"

    report += "ПРОБЛЕМНЫЕ МЕСТА:\n\n"

    if analysis.result.get('issues'):
        for issue in analysis.result['issues']:
            report += f"• {issue.get('text', '')}\n"
            report += f"  - {issue.get('explanation', '')}\n\n"
    else:
        report += "Проблемы не обнаружены.\n\n"

    report += "РЕКОМЕНДАЦИИ:\n\n"

    if analysis.result.get('suggestions'):
        for suggestion in analysis.result['suggestions']:
            report += f"• {suggestion}\n"
    else:
        report += "Рекомендации отсутствуют.\n"

    # Создаем файл
    response = HttpResponse(content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="report_{analysis.id}.txt"'
    response.write(report)
    return response
@login_required
def download_word(request, analysis_id):
    analysis = Analysis.objects.get(id=analysis_id, user=request.user)
    doc = DocxDocument()

    # Заголовок
    doc.add_heading('Отчет по анализу документа', level=1)

    # Проблемы
    doc.add_heading('Выявленные риски:', level=2)
    for issue in analysis.result.get('issues', []):
        doc.add_paragraph(f"⚠️ {issue['text']}", style='ListBullet')

    # Рекомендации
    doc.add_heading('Рекомендации:', level=2)
    for suggestion in analysis.result.get('suggestions', []):
        doc.add_paragraph(f"• {suggestion}", style='ListBullet')

    # Сохранение в буфер
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer,
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="report_{analysis.id}.docx"'
    return response